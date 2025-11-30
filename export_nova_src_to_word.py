#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Nova RH – Export de tout le code source .ts et .tsx vers un document Word
-------------------------------------------------------------------------
Ce script lit récursivement tous les fichiers du dossier /src dans ton projet
(localisé ici : ~/Téléchargements/novavfcopie3/) et compile le code dans un .docx.
"""

import os
from docx import Document
from pathlib import Path

# 📁 Chemin absolu de ton projet
HOME = Path.home()
SRC_DIR = HOME / "desktop" / "novavfcopie3" / "src"
OUTPUT_FILE = HOME / "downloads" / "NovaRH_Source_Code.docx"

# 🧱 Création du document Word
doc = Document()
doc.add_heading("Nova RH – Export du code source TypeScript / TSX", level=1)
doc.add_paragraph(
    "Ce document contient l’intégralité du code source des fichiers .ts et .tsx "
    "du projet Nova RH (dossier /src). Chaque section affiche le chemin du fichier "
    "et son contenu complet pour analyse technique."
)

# 🔧 Fonction utilitaire pour lire un fichier
def read_file(filepath):
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"[❌ Erreur lecture fichier : {e}]"

# 🔍 Parcours récursif du dossier /src
if not SRC_DIR.exists():
    print(f"❌ Le dossier {SRC_DIR} n'existe pas.")
    exit()

for root, _, files in os.walk(SRC_DIR):
    for file in files:
        if file.endswith(".ts") or file.endswith(".tsx"):
            full_path = os.path.join(root, file)
            rel_path = os.path.relpath(full_path, SRC_DIR)

            # Titre du fichier
            doc.add_heading(f"📄 {rel_path}", level=2)

            # Lecture du contenu
            content = read_file(full_path)
            doc.add_paragraph(content)

# 💾 Sauvegarde du document Word
doc.save(OUTPUT_FILE)
print(f"✅ Export terminé ! Fichier généré : {OUTPUT_FILE}")
